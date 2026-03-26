"""Student-facing API endpoints."""

import json
from fastapi import APIRouter, Depends, Request
from fastapi.responses import StreamingResponse
from auth import get_current_user, CurrentUser
from models.schemas import (
    CoverageResponse,
    SyllabusResponse,
    ReadinessResponse,
    DoubtRequest,
    DigestResponse,
    ForgeResponse,
    QuestionListResponse,
    MockGenerateResponse,
    MockAsyncResponse,
    IngestPaperRequest,
    IngestPaperResponse,
    PatternMinerResponse,
    TutorChatRequest,
    WeakSpotsResponse,
    MockSubmitRequest,
    MockSubmitResponse,
    RevisionGenerateRequest,
    RevisionScheduleResponse,
    RebalanceRequest,
    TodoListResponse,
    UpdateTodoRequest,
    ResearchRequest,
    ExamStartResponse,
    AutosaveRequest,
    SubmitExamRequest,
    IntegrityEventRequest,
    StudyTimePredictorResponse,
)
from engines.ace.chaos_cleaner import chaos_cleaner
from engines.ace.syllabus_mapper import syllabus_mapper
from engines.lre.doc_doubt import doc_doubt
from engines.ace.lecture_digest import lecture_digest
from engines.ace.question_forge import question_forge
from engines.ape.smart_mock import smart_mock
from engines.ace.paper_pattern_miner import paper_pattern_miner
from engines.lre.memory_tutor import memory_tutor
from engines.ape.weak_spotter import weak_spotter
from engines.ape.revision_clock import revision_clock
from engines.ape.study_time_predictor import study_time_predictor

router = APIRouter()


# ─── Coverage & Syllabus ───────────────────────────────────

@router.get("/coverage", response_model=CoverageResponse)
async def get_coverage(user: CurrentUser = Depends(get_current_user)):
    """Get syllabus coverage map for the authenticated student."""
    coverage = await chaos_cleaner.get_coverage(
        user_id=user.user_id,
        institution_id=user.institution_id or "",
    )
    return CoverageResponse(coverage=coverage)


@router.get("/syllabus", response_model=SyllabusResponse | None)
async def get_syllabus(user: CurrentUser = Depends(get_current_user)):
    """Get the parsed syllabus for the student's institution."""
    syllabus = await syllabus_mapper.get_syllabus(
        institution_id=user.institution_id or "",
    )
    if not syllabus:
        return None
    return SyllabusResponse(**syllabus)


# ─── DocDoubt — RAG Chat ───────────────────────────────────

@router.post("/doubt")
async def doubt_chat(body: DoubtRequest, user: CurrentUser = Depends(get_current_user)):
    """Stream a grounded answer to a student's question via SSE.

    Uses DocDoubt to search uploaded notes + faculty courseware,
    then streams the LLM response with inline citations.
    """

    async def event_stream():
        # Send source citations first so the frontend can display them
        sources = await doc_doubt.get_sources(
            question=body.message,
            user_id=user.user_id,
            institution_id=user.institution_id or "",
        )
        yield f"data: {json.dumps({'type': 'sources', 'data': sources})}\n\n"

        # Stream the response text
        full_response = ""
        async for chunk in doc_doubt.chat(
            question=body.message,
            user_id=user.user_id,
            institution_id=user.institution_id or "",
            conversation_history=None,
        ):
            full_response += chunk
            yield f"data: {json.dumps({'type': 'text', 'data': chunk})}\n\n"

        # Signal completion
        yield f"data: {json.dumps({'type': 'done', 'data': full_response})}\n\n"

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@router.get("/doubt/history")
async def doubt_history(user: CurrentUser = Depends(get_current_user)):
    """List past DocDoubt conversations for the student.

    Stub — will query Supabase chat_history table.
    """
    # TODO: Query Supabase for user's chat_history where agent='doc_doubt'
    return {"conversations": []}


# ─── LectureDigest ─────────────────────────────────────────

@router.post("/digest/{doc_id}", response_model=DigestResponse)
async def generate_digest(
    doc_id: str,
    user: CurrentUser = Depends(get_current_user),
):
    """Generate a structured summary for a specific uploaded document.

    Returns key concepts, definitions, formulas, and review questions.
    """
    result = await lecture_digest.digest(
        doc_id=doc_id,
        user_id=user.user_id,
        institution_id=user.institution_id or "",
    )
    return DigestResponse(**result)


# ─── QuestionForge ─────────────────────────────────────────

@router.post("/forge/{doc_id}", response_model=ForgeResponse)
async def forge_questions(
    doc_id: str,
    user: CurrentUser = Depends(get_current_user),
):
    """Generate a Bloom's-tagged question bank from an uploaded document.

    Scrolls the document's Qdrant chunks in batches and calls Haiku
    to produce 5 questions per batch spanning all Bloom's levels.
    Questions are persisted to the question_bank table.
    """
    result = await question_forge.forge_from_document(
        doc_id=doc_id,
        user_id=user.user_id,
        institution_id=user.institution_id or "",
    )
    return ForgeResponse(**result)


@router.get("/questions", response_model=QuestionListResponse)
async def list_questions(user: CurrentUser = Depends(get_current_user)):
    """List all questions in the student's question bank.

    Stub — will query Supabase question_bank once SmartMock is built.
    """
    # TODO: SELECT * FROM question_bank WHERE user_id = user.user_id
    return QuestionListResponse(questions=[], total=0)


# ─── SmartMock ─────────────────────────────────────────────

@router.post("/mock/generate", response_model=MockGenerateResponse)
async def generate_mock(user: CurrentUser = Depends(get_current_user)):
    """Generate a personalised mock exam paper for the student.

    Selects questions from the question bank, overweighting weak topics
    (2× probability). Renders a PDF via reportlab, uploads to R2, and
    persists a mock record in Supabase.
    """
    result = await smart_mock.generate(
        user_id=user.user_id,
        institution_id=user.institution_id or "",
    )
    return MockGenerateResponse(**result)


@router.post("/mock/{mock_id}/submit", response_model=MockSubmitResponse)
async def submit_mock(
    mock_id: str,
    body: MockSubmitRequest,
    user: CurrentUser = Depends(get_current_user),
):
    """Grade a mock submission and update weak spots.

    (Stub implementation — assumes a separate grading step computes scores,
    then we pass those scores to the WeakSpotter.)
    """
    # STUB: Mocked grading logic
    stub_topic_breakdown = [
        {"topic": "Quantum States", "unit": 1, "score": 4.0, "max_score": 10.0},
        {"topic": "Entanglement", "unit": 2, "score": 9.0, "max_score": 10.0},
    ]

    await weak_spotter.update_from_mock(user.user_id, stub_topic_breakdown)

    return MockSubmitResponse(
        score=13.0,
        max_score=20.0,
        topic_breakdown=stub_topic_breakdown,
    )


# ─── MemoryTutor ───────────────────────────────────────────

@router.post("/tutor")
async def chat_with_tutor(
    body: TutorChatRequest,
    user: CurrentUser = Depends(get_current_user),
):
    """Chat with the persistent MemoryTutor.

    Streams a Sonnet response injected with a briefing of the student's
    prior sessions and known weak spots.
    """
    conversation = [m.model_dump() for m in body.conversation]

    return StreamingResponse(
        memory_tutor.chat(
            message=body.message,
            user_id=user.user_id,
            institution_id=user.institution_id or "",
            conversation=conversation,
        ),
        media_type="text/event-stream",
    )


# ─── WeakSpotter ───────────────────────────────────────────

@router.get("/weak-spots", response_model=WeakSpotsResponse)
async def get_weak_spots(user: CurrentUser = Depends(get_current_user)):
    """Retrieve the student's concept gap tracker, ordered weakest first."""
    spots = await weak_spotter.get_weak_spots(user.user_id)
    return WeakSpotsResponse(weak_spots=spots)


# ─── RevisionClock ─────────────────────────────────────────

@router.post("/schedule", response_model=RevisionScheduleResponse)
async def generate_schedule(
    body: RevisionGenerateRequest,
    user: CurrentUser = Depends(get_current_user),
):
    """Generate a day-by-day revision schedule up to the exam date."""
    schedule = await revision_clock.generate(
        user_id=user.user_id,
        institution_id=user.institution_id or "",
        exam_date=body.exam_date,
    )
    return RevisionScheduleResponse(
        schedule=schedule,
        exam_date=str(body.exam_date),
    )


@router.get("/schedule", response_model=RevisionScheduleResponse)
async def get_schedule(
    user: CurrentUser = Depends(get_current_user),
):
    """Retrieve the student's existing total revision schedule."""
    data = await revision_clock.get_schedule(user.user_id)
    return RevisionScheduleResponse(
        schedule=data.get("schedule", []),
        exam_date=data.get("exam_date", ""),
    )


@router.post("/schedule/rebalance", response_model=RevisionScheduleResponse)
async def rebalance_schedule(
    body: RebalanceRequest,
    user: CurrentUser = Depends(get_current_user),
):
    """Redistribute incomplete tasks from missed days across future days."""
    schedule = await revision_clock.rebalance(
        user.user_id,
        body.missed_dates,
    )
    # Re-fetch to get exam_date
    data = await revision_clock.get_schedule(user.user_id)
    return RevisionScheduleResponse(
        schedule=schedule,
        exam_date=data.get("exam_date", ""),
    )


@router.get("/todos", response_model=TodoListResponse)
async def get_today_todos(
    user: CurrentUser = Depends(get_current_user),
):
    """Get the dynamic daily to-do list derived from the master schedule.
    Automatically seeds today's rows if it's a new day.
    """
    result = await revision_clock.get_today_todos(user.user_id)
    return TodoListResponse(**result)


@router.patch("/todos", response_model=TodoListResponse)
async def update_todo_status(
    body: UpdateTodoRequest,
    user: CurrentUser = Depends(get_current_user),
):
    """Mark a daily to-do item as done, skipped, or pending."""
    result = await revision_clock.update_todo_status(
        user.user_id,
        body.task_key,
        body.status,
    )
    return TodoListResponse(**result)


# ─── PaperPatternMiner ─────────────────────────────────────

@router.get("/patterns", response_model=PatternMinerResponse)
async def get_patterns(user: CurrentUser = Depends(get_current_user)):
    """Get exam pattern insights for the student's institution.

    Aggregates historically ingested past-papers and returns topic
    frequency, unit mark distribution, question-type breakdown, and
    year-wise trends.
    """
    result = await paper_pattern_miner.get_patterns(
        institution_id=user.institution_id or "",
    )
    return PatternMinerResponse(**result)


@router.post("/patterns/ingest", response_model=IngestPaperResponse)
async def ingest_paper(
    body: IngestPaperRequest,
    user: CurrentUser = Depends(get_current_user),
):
    """Ingest a past exam paper and extract topic-tagged questions.

    Uses Haiku to extract every question, matches each to a syllabus
    topic via Qdrant, and persists the enriched questions to Supabase.
    """
    result = await paper_pattern_miner.ingest_paper(
        paper_text=body.paper_text,
        year=body.year,
        institution_id=user.institution_id or "",
    )
    return IngestPaperResponse(**result)


# ─── Readiness ─────────────────────────────────────────────

@router.get("/readiness", response_model=ReadinessResponse)
async def get_readiness(user: CurrentUser = Depends(get_current_user)):
    """Compute real-time readiness score (Section 9 formula).

    score = mock_avg*0.40 + gap_coverage*0.25 + consistency*0.20
          + syllabus_coverage*0.15 - urgency_penalty
    """
    from engines.ape.readiness import readiness_engine
    result = await readiness_engine.compute(
        user_id=user.user_id,
        institution_id=user.institution_id or "",
    )
    return ReadinessResponse(**result)


# ─── AutoResearcher ─────────────────────────────────────────

@router.post("/research")
async def run_research(
    body: ResearchRequest,
    user: CurrentUser = Depends(get_current_user),
):
    """5-stage multi-agent report builder (SSE stream).

    Stages: Planner → Researcher → Writer → Editor → Presenter
    Each stage emits an SSE event. Final event contains the full report
    and slide outline.
    """
    from engines.lre.auto_researcher import auto_researcher

    return StreamingResponse(
        auto_researcher.run(topic=body.topic, user_id=user.user_id),
        media_type="text/event-stream",
    )


# ─── AutoResearcher report download (DOCX) ────────────────────────────

@router.get("/research/{report_id}/download")
async def download_research_report(
    report_id: str,
    user: CurrentUser = Depends(get_current_user),
):
    """Download an AutoResearcher report as a .docx file.

    Fetches the saved Markdown report from Supabase, converts it to
    a Word document via python-docx, and streams the file back.
    """
    from fastapi import HTTPException
    from fastapi.responses import Response
    from engines.lre.auto_researcher import auto_researcher

    report = await auto_researcher.get_report(report_id, user.user_id)
    if not report:
        raise HTTPException(status_code=404, detail="Report not found.")

    docx_bytes = _markdown_to_docx(report["report_md"], report.get("topic", "Research Report"))
    filename = f"bluescholar_report_{report_id[:8]}.docx"
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _markdown_to_docx(markdown_text: str, title: str) -> bytes:
    """Convert Markdown text to DOCX bytes using python-docx."""
    from io import BytesIO
    try:
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        doc.core_properties.title = title

        # Title
        doc.add_heading(title, level=0)

        for line in markdown_text.splitlines():
            stripped = line.strip()
            if stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith(("- ", "* ")):
                p = doc.add_paragraph(stripped[2:], style="List Bullet")
            elif stripped:
                doc.add_paragraph(stripped)
            else:
                doc.add_paragraph("")

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception:
        return b""  # fallback: empty bytes


# ─── ExamArena (student-side) ───────────────────────────────────

@router.get("/exams/active")
async def list_active_exams(user: CurrentUser = Depends(get_current_user)):
    """List exams currently open for the student's institution."""
    try:
        from config import get_settings
        from supabase import create_client
        settings = get_settings()
        sb = create_client(settings.supabase_url, settings.supabase_service_role_key)
        rows = (
            sb.table("exams")
            .select("id, title, time_limit_mins, opens_at, closes_at, instructions")
            .eq("institution_id", user.institution_id)
            .eq("status", "open")
            .execute()
        )
        return {"exams": rows.data or []}
    except Exception:
        return {"exams": []}


@router.post("/exams/{exam_id}/start", response_model=ExamStartResponse)
async def start_exam(
    exam_id: str,
    user: CurrentUser = Depends(get_current_user),
):
    """Create or return the student's submission row and start the timer."""
    import uuid
    from datetime import datetime, timezone
    started_at = datetime.now(timezone.utc).isoformat()
    submission_id = str(uuid.uuid4())
    try:
        from config import get_settings
        from supabase import create_client
        settings = get_settings()
        sb = create_client(settings.supabase_url, settings.supabase_service_role_key)
        # Upsert so re-starting doesn't create duplicate rows
        res = sb.table("exam_submissions").upsert({
            "id": submission_id,
            "exam_id": exam_id,
            "student_id": user.user_id,
            "answers": {},
            "started_at": started_at,
            "grading_status": "pending",
        }).execute()
        if res.data:
            submission_id = res.data[0].get("id", submission_id)
            started_at = res.data[0].get("started_at", started_at)
    except Exception:
        pass
    return ExamStartResponse(
        submission_id=submission_id,
        exam_id=exam_id,
        started_at=started_at,
    )


@router.put("/exams/{exam_id}/autosave")
async def autosave_exam(
    exam_id: str,
    body: AutosaveRequest,
    user: CurrentUser = Depends(get_current_user),
):
    """Debounced answer auto-save (frontend calls every 30 s)."""
    try:
        from config import get_settings
        from supabase import create_client
        settings = get_settings()
        sb = create_client(settings.supabase_url, settings.supabase_service_role_key)
        (
            sb.table("exam_submissions")
            .update({"answers": body.answers})
            .eq("exam_id", exam_id)
            .eq("student_id", user.user_id)
            .execute()
        )
    except Exception:
        pass
    return {"status": "saved"}


@router.post("/exams/{exam_id}/submit")
async def submit_exam(
    exam_id: str,
    body: SubmitExamRequest,
    user: CurrentUser = Depends(get_current_user),
):
    """Final exam submission. Saves answers and marks as submitted."""
    from datetime import datetime, timezone
    submitted_at = datetime.now(timezone.utc).isoformat()
    try:
        from config import get_settings
        from supabase import create_client
        settings = get_settings()
        sb = create_client(settings.supabase_url, settings.supabase_service_role_key)
        (
            sb.table("exam_submissions")
            .update({"answers": body.answers, "submitted_at": submitted_at})
            .eq("exam_id", exam_id)
            .eq("student_id", user.user_id)
            .execute()
        )
    except Exception:
        pass
    return {"status": "submitted", "submitted_at": submitted_at}


@router.post("/exams/integrity")
async def log_integrity_events(
    body: IntegrityEventRequest,
    user: CurrentUser = Depends(get_current_user),
):
    """GuardEye: batch-upload browser integrity events (tab switches, copy attempts, etc.)."""
    try:
        from config import get_settings
        from supabase import create_client
        settings = get_settings()
        sb = create_client(settings.supabase_url, settings.supabase_service_role_key)
        (
            sb.table("exam_submissions")
            .update({"integrity_flags": body.events})
            .eq("exam_id", body.exam_id)
            .eq("student_id", body.student_id)
            .execute()
        )
    except Exception:
        pass
    return {"status": "logged", "event_count": len(body.events)}


# ─── Past Research Reports ──────────────────────────────────

@router.get("/research")
async def list_research_reports(user: CurrentUser = Depends(get_current_user)):
    """List past AutoResearcher reports for the student."""
    try:
        from config import get_settings
        from supabase import create_client
        settings = get_settings()
        sb = create_client(settings.supabase_url, settings.supabase_service_role_key)
        rows = (
            sb.table("research_reports")
            .select("id, topic, created_at")
            .eq("user_id", user.user_id)
            .order("created_at", desc=True)
            .execute()
        )
        return {"reports": rows.data or []}
    except Exception:
        return {"reports": []}


# ─── Mock Retrieval ─────────────────────────────────────────

@router.get("/mock/{mock_id}")
async def get_mock(
    mock_id: str,
    user: CurrentUser = Depends(get_current_user),
):
    """Retrieve a specific mock paper by ID."""
    try:
        from config import get_settings
        from supabase import create_client
        settings = get_settings()
        sb = create_client(settings.supabase_url, settings.supabase_service_role_key)
        row = (
            sb.table("mocks")
            .select("*")
            .eq("id", mock_id)
            .eq("user_id", user.user_id)
            .limit(1)
            .execute()
        )
        return row.data[0] if row.data else {"error": "Mock not found"}
    except Exception:
        return {"error": "Could not fetch mock"}


# ─── Async Mock Generation ─────────────────────────────────────────

@router.post("/mock/generate-async", response_model=MockAsyncResponse)
async def generate_mock_async(
    user: CurrentUser = Depends(get_current_user),
):
    """Enqueue mock paper generation as a Celery task.

    Returns immediately with a task ID. Poll GET /student/mock/{mock_id}
    to check when the paper is ready (status will change from 'pending').
    Prefer this endpoint for larger question banks where generation may take
    several seconds.
    """
    from workers.tasks.mock_gen import generate_mock_paper
    task = generate_mock_paper.delay(
        user_id=user.user_id,
        institution_id=user.institution_id or "",
    )
    return MockAsyncResponse(task_id=task.id, status="queued")


# ─── StudyTimePredictor ──────────────────────────────────────────────

@router.get("/study-windows", response_model=StudyTimePredictorResponse)
async def get_study_windows(
    user: CurrentUser = Depends(get_current_user),
):
    """Return the student's predicted peak study hours.

    Analyses chat, todo, and mock timestamps to surface the hours of day
    when the student is most active. Used by RevisionClock for scheduling.
    """
    result = await study_time_predictor.predict(user_id=user.user_id)
    # Convert int keys to str for Pydantic dict[str, int]
    result["hour_distribution"] = {str(k): v for k, v in result["hour_distribution"].items()}
    return StudyTimePredictorResponse(**result)
